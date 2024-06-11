# coding=utf-8
import sys
import os
import pandas as pd
from copy import deepcopy
from docx import Document
from string import Template

# 资源文件打包到exe时，生成资源文件的访问路径
def resource_path(relative_path):
    if getattr(sys, 'frozen', False): #是否Bundle Resource
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# 清理字符串，去除换行符
def clean_string(input_string):
    if isinstance(input_string, str):
        return input_string.replace('\n', '')
    else:
        return input_string

def add_table(docx_obj, mb_obj, data_obj): 

# 数据对象，用于填充模板 
    data = {
        'dybs'  : data_obj['单元标识'],
        'sjzs'  : data_obj['设计追踪'],
        'gnms'  : data_obj['功能描述'],
        'qdmk'  : data_obj['驱动模块'],
        'dzmk'  : clean_string(data_obj['打桩模块']),
        'fgljl' : clean_string(data_obj['覆盖率记录']),
        'bz'    : data_obj['备注'],
        'csylmc': data_obj['测试用例名称'],
        'csylbs': data_obj['测试用例标识'],
        'cslx'  : data_obj['测试类型'],
        'hdz'   : data_obj['活动桩'],
        'cryj'  :  '\n' + data_obj['插入语句'] if '\n' in data_obj['插入语句'] else data_obj['插入语句'] ,
        'qjbl'  :  '\n' + data_obj['创建全局变量'] if '\n' in data_obj['创建全局变量'] else data_obj['创建全局变量'] ,
        'cssm'  : data_obj['测试说明'],
        'srblm' : data_obj['输入变量名称'],
        'srblqz': data_obj['取值'],
        'scblm' : data_obj['输出变量名称'],
        'yqsc'  : data_obj['预期输出'],
        'sjsc'  : data_obj['实际输出'],
        'sjz'   : data_obj['测试者/校对者'],
        'xdz'   : data_obj['测试者/校对者'],
        'csry'  : data_obj['测试者/校对者'],
        'cszxsj': data_obj['测试时间'].date() if isinstance(data_obj['测试时间'], pd.Timestamp) else data_obj['测试时间'],
    }
  
# 如果测试者/校对者字段包含'/'分隔符，则分别赋值给相应的变量
    cs_jz = str(data_obj['测试者/校对者'])
    if '/' in cs_jz:
        data['sjz'] = cs_jz.split('/')[0]
        data['xdz'] = cs_jz.split('/')[1]
        data['csry'] = cs_jz.split('/')[0]
    else:
        data['sjz'] = cs_jz
        data['xdz'] = cs_jz
        data['csry'] = cs_jz

# 复制模板中的表格
    new_table = deepcopy(mb_obj.tables[0])

# 填充表格中的文本内容    
    for row in new_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run_template = Template(run.text)
                    run.text = run_template.substitute(data)

# 在文档中添加新表格                        
    paragraph = docx_obj.add_paragraph(f"{data_obj['测试用例名称']}")
    paragraph._p.addnext(new_table._element)
    docx_obj.add_page_break()

if __name__ == '__main__':
    if len(sys.argv) != 2:                              # 检查命令行参数是否正确
        print('\033[0;31mneed excel file\033[0m')
        os.system("pause")
        sys.exit(1)
    excel_file = sys.argv[1]
    # excel_file = 'dist/单元测试协同表格.xlsx'
    excel_datas = pd.read_excel(excel_file,sheet_name=None)             # 读取Excel文件中的数据，读取后对 DataFrame 进行转换
    excel_datas = {sheet_name: df.astype(str) for sheet_name, df in excel_datas.items()}

    template_docx = Document(resource_path("template/template.docx"))    # 加载Word文档模板
    new_docx = Document(resource_path("template/new.docx"))   # 创建新的Word文档对象

    template_count = 0  # 初始化模板计数器
    
    for sheet_name, df in excel_datas.items():                           # 遍历Excel数据，生成测试文档
        if sheet_name in ['test', '索引']:
            continue
        for index, row in df.iterrows():
            print(sheet_name, row['测试用例名称'])
            add_table(new_docx,template_docx,row)
            template_count += 1  # 每成功填充一次模板，计数器加一

    save_file_name = os.path.splitext(excel_file)[0] + '_oa.docx'
    new_docx.save(save_file_name)
    
    print(f"\033[0;32m{save_file_name} is created\033[0m")
    print(f"\033[0;32m一共执行了{template_count}个测试用例\033[0m")  # 打印生成的用例次数

    os.system("pause")

# ☑☐□